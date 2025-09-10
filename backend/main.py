from fastapi import FastAPI, File, UploadFile, Form, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import os
import uuid
import tempfile
import io
from typing import List, Dict, Any, Optional
import logging
import re
from datetime import datetime
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
from dotenv import load_dotenv

# Load environment variables from .env file
load_dotenv()

# Document processing imports
import pdfplumber
from docx import Document
try:
    import pytesseract
    from PIL import Image
    import fitz  # PyMuPDF
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

# LangChain imports
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings, OpenAI, ChatOpenAI
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain.docstore.document import Document as LangChainDocument
from langchain_core.prompts import ChatPromptTemplate, PromptTemplate
from langchain.memory import ConversationBufferWindowMemory
from langchain.chains import ConversationalRetrievalChain
from langchain.schema import BaseMessage, HumanMessage, AIMessage

# Pinecone import
from pinecone import Pinecone

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Log OCR availability after logger is initialized
if not OCR_AVAILABLE:
    logger.warning("OCR libraries not available. Image-only PDFs will not be processed.")

# Initialize FastAPI app
app = FastAPI(title="Multi-Step Reasoning AI Assistant", version="1.0.0")

# Add CORS middleware for React frontend
app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:3000"],  # React dev server
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Environment variables
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")

# Email configuration (optional)
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
EMAIL_ADDRESS = os.getenv("EMAIL_ADDRESS")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

if not OPENAI_API_KEY:
    raise ValueError("OPENAI_API_KEY environment variable is required")
if not PINECONE_API_KEY:
    raise ValueError("PINECONE_API_KEY environment variable is required")

# Log email configuration status
if EMAIL_ADDRESS and EMAIL_PASSWORD:
    logger.info("Email configuration found - email reports will be sent")
else:
    logger.warning("Email configuration missing - email reports will be logged only")

# Initialize Pinecone
pc = Pinecone(api_key=os.getenv("PINECONE_API_KEY"))
index_name = "ai-assitant"  # Using your specific index name
index = pc.Index(index_name, host="https://ai-assitant-zlht6dl.svc.aped-4627-b74a.pinecone.io")

# Initialize LangChain components
# Initialize OpenAI embeddings with text-embedding-3-small model (512 dimensions)
embeddings = OpenAIEmbeddings(
    openai_api_key=OPENAI_API_KEY,
    model="text-embedding-3-small",
    dimensions=512  # Explicitly set dimensions for text-embedding-3-small
)
llm = ChatOpenAI(temperature=0.7, openai_api_key=OPENAI_API_KEY, model="gpt-3.5-turbo")

# Initialize conversation memory storage (in-memory for now, can be moved to persistent storage)
user_memories = {}

# Enhanced AI Assistant Prompt Template
enhanced_system_prompt = """
You are an intelligent and helpful AI assistant with access to the user's uploaded documents. You have a friendly, professional personality and excellent memory of previous conversations.

Key Instructions:
1. Always provide helpful, accurate, and contextual responses based on the available documents
2. Remember previous conversations and refer to them when relevant
3. If information is missing or unclear in the documents, explicitly state what is available vs. unavailable
4. Be conversational and personable while maintaining professionalism
5. When greeting users, use their name if you know it
6. Provide specific details from documents rather than generic responses
7. If asked about something not in the documents, clearly state that and offer to help with what is available

Document Context:
{context}

Conversation History:
{chat_history}

Current Question: {question}

Provide a helpful and contextual response:"""

# Initialize QA chain with basic prompt (will be enhanced per conversation)
qa_prompt = ChatPromptTemplate.from_template("{context}")
qa_chain = create_stuff_documents_chain(llm, qa_prompt)

# Text splitter for chunking documents
text_splitter = RecursiveCharacterTextSplitter(
    chunk_size=500,
    chunk_overlap=50,
    length_function=len,
)

# Pydantic models
class QuestionRequest(BaseModel):
    question: str
    guest_id: str

class AnswerResponse(BaseModel):
    answer: str
    sources: List[Dict[str, Any]]

class UploadResponse(BaseModel):
    message: str
    document_name: str
    chunks_processed: int

class UserProfile(BaseModel):
    name: str
    email: str = None
    guest_id: str

class UserProfileResponse(BaseModel):
    message: str
    profile: Dict[str, Any]

class EmailReportRequest(BaseModel):
    guest_id: str
    email: str
    report_type: str = "conversation_summary"

class GenerateReportRequest(BaseModel):
    guest_id: str
    report_topic: str
    include_email: bool = False
    email_address: str = None

class DocumentStats(BaseModel):
    total_chunks: int
    total_documents: int
    documents: List[Dict[str, Any]]
    warning_message: Optional[str] = None

class DeleteDocumentRequest(BaseModel):
    guest_id: str
    document_name: str

class ConversationMessage(BaseModel):
    role: str  # 'human' or 'ai'
    content: str
    timestamp: str

# Helper functions for conversation memory management
def get_user_memory(guest_id: str) -> ConversationBufferWindowMemory:
    """Get or create conversation memory for a user"""
    if guest_id not in user_memories:
        user_memories[guest_id] = ConversationBufferWindowMemory(
            k=10,  # Keep last 10 exchanges
            return_messages=True,
            memory_key="chat_history"
        )
    return user_memories[guest_id]

def store_conversation_in_pinecone(guest_id: str, question: str, answer: str):
    """Store conversation exchange in Pinecone for persistent memory"""
    try:
        conversation_text = f"Human: {question}\nAI: {answer}"
        conversation_embedding = embeddings.embed_query(conversation_text)
        
        conversation_id = f"conversation_{guest_id}_{uuid.uuid4().hex[:8]}"
        
        index.upsert(vectors=[{
            "id": conversation_id,
            "values": conversation_embedding,
            "metadata": {
                "type": "conversation",
                "user_id": guest_id,
                "question": question,
                "answer": answer,
                "timestamp": datetime.now().isoformat(),
                "conversation_text": conversation_text
            }
        }])
        
        logger.info(f"Stored conversation for user {guest_id}")
    except Exception as e:
        logger.error(f"Error storing conversation: {e}")

def get_relevant_conversation_history(guest_id: str, current_question: str, top_k: int = 3) -> str:
    """Retrieve relevant past conversations for context"""
    try:
        question_embedding = embeddings.embed_query(current_question)
        
        # Query for relevant past conversations
        query_response = index.query(
            vector=question_embedding,
            top_k=top_k,
            include_metadata=True,
            filter={"user_id": guest_id, "type": "conversation"}
        )
        
        if query_response.matches:
            relevant_history = []
            for match in query_response.matches:
                if match.score > 0.7:  # Only include highly relevant conversations
                    relevant_history.append(match.metadata.get('conversation_text', ''))
            
            return "\n\n".join(relevant_history) if relevant_history else ""
        
        return ""
    except Exception as e:
        logger.error(f"Error retrieving conversation history: {e}")
        return ""

# Helper functions
def send_email(to_email: str, subject: str, body: str) -> bool:
    """Send email using SMTP configuration"""
    if not EMAIL_ADDRESS or not EMAIL_PASSWORD:
        logger.warning("Email configuration missing - cannot send email")
        return False
    
    try:
        # Create message
        msg = MIMEMultipart()
        msg['From'] = EMAIL_ADDRESS
        msg['To'] = to_email
        msg['Subject'] = subject
        
        # Add body to email
        msg.attach(MIMEText(body, 'plain'))
        
        # Create SMTP session
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()  # Enable security
        server.login(EMAIL_ADDRESS, EMAIL_PASSWORD)
        
        # Send email
        text = msg.as_string()
        server.sendmail(EMAIL_ADDRESS, to_email, text)
        server.quit()
        
        logger.info(f"Email sent successfully to {to_email}")
        return True
        
    except Exception as e:
        logger.error(f"Failed to send email to {to_email}: {e}")
        return False

def extract_text_from_pdf(file_path: str) -> str:
    """Extract text from PDF using pdfplumber, with OCR fallback for image-only PDFs"""
    text = ""
    try:
        # First, try to extract text using pdfplumber
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        
        # If no text was extracted and OCR is available, try OCR
        if not text.strip() and OCR_AVAILABLE:
            logger.info("No text found in PDF, attempting OCR extraction")
            text = extract_text_with_ocr(file_path)
        
        # If still no text, accept the PDF but note it's image-only
        if not text.strip():
            logger.info("PDF appears to contain only images and OCR is not available")
            text = "[This document contains images that could not be processed for text extraction. OCR capabilities are not available.]"
            
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing PDF: {str(e)}")
    return text

def extract_text_with_ocr(file_path: str) -> str:
    """Extract text from PDF using OCR (for image-only PDFs)"""
    if not OCR_AVAILABLE:
        return ""
    
    text = ""
    try:
        # Open PDF with PyMuPDF
        pdf_document = fitz.open(file_path)
        
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            
            # Convert page to image
            pix = page.get_pixmap()
            img_data = pix.tobytes("ppm")
            
            # Convert to PIL Image
            img = Image.open(io.BytesIO(img_data))
            
            # Extract text using OCR
            page_text = pytesseract.image_to_string(img)
            if page_text.strip():
                text += page_text + "\n"
        
        pdf_document.close()
        
        if text.strip():
            logger.info(f"Successfully extracted text using OCR: {len(text)} characters")
        
    except Exception as e:
        logger.error(f"Error during OCR extraction: {e}")
        text = "[This document contains images that could not be processed due to OCR error.]"
    
    return text

def extract_text_from_docx(file_path: str) -> str:
    """Extract text from DOCX using python-docx"""
    text = ""
    try:
        doc = Document(file_path)
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing DOCX: {str(e)}")
    return text

def extract_text_from_txt(file_path: str) -> str:
    """Extract text from TXT file"""
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    except Exception as e:
        logger.error(f"Error extracting text from TXT: {e}")
        raise HTTPException(status_code=400, detail=f"Error processing TXT: {str(e)}")

def process_document(file_path: str, filename: str) -> str:
    """Process document based on file extension"""
    file_extension = filename.lower().split('.')[-1]
    
    if file_extension == 'pdf':
        return extract_text_from_pdf(file_path)
    elif file_extension == 'docx':
        return extract_text_from_docx(file_path)
    elif file_extension == 'txt':
        return extract_text_from_txt(file_path)
    else:
        raise HTTPException(status_code=400, detail=f"Unsupported file type: {file_extension}")

@app.get("/")
async def root():
    return {"message": "Multi-Step Reasoning AI Assistant API"}

@app.get("/health")
async def health_check():
    return {"status": "healthy"}

@app.get("/documents/{guest_id}", response_model=DocumentStats)
async def get_document_stats(guest_id: str):
    """Get document statistics for a guest user"""
    try:
        # Query all documents for this user
        doc_response = index.query(
            vector=[0] * 512,
            top_k=1000,  # Get all chunks
            include_metadata=True,
            filter={"user_id": guest_id, "type": {"$ne": "user_profile"}}
        )
        
        total_chunks = len(doc_response.matches)
        
        # Group by document name
        doc_groups = {}
        for match in doc_response.matches:
            doc_name = match.metadata.get('document_name')
            # Skip chunks without proper document names
            if not doc_name or doc_name == 'Unknown':
                continue
            if doc_name not in doc_groups:
                doc_groups[doc_name] = {
                    'name': doc_name,
                    'chunks': 0,
                    'upload_date': match.metadata.get('upload_date', 'Unknown')
                }
            doc_groups[doc_name]['chunks'] += 1
        
        documents = list(doc_groups.values())
        warning_message = None
        
        if total_chunks > 10:
            warning_message = f"You have {total_chunks} chunks uploaded. Consider deleting old documents to improve performance."
        
        return DocumentStats(
            total_chunks=total_chunks,
            total_documents=len(documents),
            documents=documents,
            warning_message=warning_message
        )
        
    except Exception as e:
        logger.error(f"Error getting document stats: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving document statistics: {str(e)}")

@app.get("/stats/{guest_id}")
async def get_user_stats(guest_id: str):
    """Get comprehensive statistics for a guest user including all data types"""
    try:
        # Query all data for this user
        all_response = index.query(
            vector=[0] * 512,
            top_k=10000,  # Get all chunks
            include_metadata=True,
            filter={"user_id": guest_id}
        )
        
        total_chunks = len(all_response.matches)
        document_chunks = 0
        conversation_chunks = 0
        other_chunks = 0
        
        # Categorize chunks by type
        for match in all_response.matches:
            chunk_type = match.metadata.get('type', 'document')
            if chunk_type == 'conversation':
                conversation_chunks += 1
            elif chunk_type == 'document' or 'document_name' in match.metadata:
                document_chunks += 1
            elif chunk_type != 'user_profile':  # Exclude user profile from counts
                other_chunks += 1
        
        return {
            "total_chunks": total_chunks,
            "document_chunks": document_chunks,
            "conversation_chunks": conversation_chunks,
            "other_chunks": other_chunks
        }
        
    except Exception as e:
        logger.error(f"Error getting user stats: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving user statistics: {str(e)}")

@app.post("/upload", response_model=UploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    guest_id: str = Form(...)
):
    """Upload and process document for a specific guest user"""
    
    # Check existing chunk count and document name uniqueness before upload
    try:
        existing_response = index.query(
            vector=[0] * 512,
            top_k=1000,
            include_metadata=True,
            filter={"user_id": guest_id, "type": {"$ne": "user_profile"}}
        )
        existing_chunks = len(existing_response.matches)
        
        # Check for document name uniqueness
        existing_documents = set()
        for match in existing_response.matches:
            if "document_name" in match.metadata:
                existing_documents.add(match.metadata["document_name"])
        
        if file.filename in existing_documents:
            raise HTTPException(
                status_code=400,
                detail=f"A document with the name '{file.filename}' already exists. Please rename your file or delete the existing document first."
            )
        
        if existing_chunks > 10:
            raise HTTPException(
                status_code=400,
                detail=f"You already have {existing_chunks} chunks uploaded. Please delete some documents before uploading new ones to maintain optimal performance."
            )
    except HTTPException:
        raise
    except Exception as e:
        logger.warning(f"Could not check existing chunks: {e}")
    
    # Validate file type
    allowed_extensions = ['pdf', 'txt', 'docx']
    file_extension = file.filename.lower().split('.')[-1]
    
    if file_extension not in allowed_extensions:
        raise HTTPException(
            status_code=400, 
            detail=f"File type not supported. Allowed types: {', '.join(allowed_extensions)}"
        )
    
    # Create temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{file_extension}") as temp_file:
        content = await file.read()
        temp_file.write(content)
        temp_file_path = temp_file.name
    
    try:
        # Extract text from document
        logger.info(f"Processing document: {file.filename} for guest: {guest_id}")
        text = process_document(temp_file_path, file.filename)
        
        if not text.strip():
            raise HTTPException(status_code=400, detail="No text content found in the document")
        
        # Split text into chunks
        chunks = text_splitter.split_text(text)
        logger.info(f"Document split into {len(chunks)} chunks")
        
        # Generate embeddings and store in Pinecone
        vectors_to_upsert = []
        
        for i, chunk in enumerate(chunks):
            # Generate embedding for chunk
            embedding = embeddings.embed_query(chunk)
            
            # Create unique ID for this chunk
            chunk_id = f"{guest_id}_{file.filename}_{i}"
            
            # Prepare metadata
            metadata = {
                "user_id": guest_id,
                "document_name": file.filename,
                "chunk_index": i,
                "text": chunk[:500],  # Store first 500 chars as snippet
                "full_text": chunk,
                "upload_date": datetime.now().isoformat()
            }
            
            vectors_to_upsert.append({
                "id": chunk_id,
                "values": embedding,
                "metadata": metadata
            })
        
        # Batch upsert to Pinecone
        batch_size = 100
        for i in range(0, len(vectors_to_upsert), batch_size):
            batch = vectors_to_upsert[i:i + batch_size]
            index.upsert(vectors=batch)
        
        logger.info(f"Successfully stored {len(chunks)} chunks for guest {guest_id}")
        
        return UploadResponse(
            message="Document uploaded and processed successfully",
            document_name=file.filename,
            chunks_processed=len(chunks)
        )
    
    except Exception as e:
        logger.error(f"Error processing document: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing document: {str(e)}")
    
    finally:
        # Clean up temporary file
         if os.path.exists(temp_file_path):
             os.unlink(temp_file_path)

@app.post("/ask", response_model=AnswerResponse)
async def ask_question(request: QuestionRequest):
    """Answer question based on uploaded documents with conversation memory"""
    
    try:
        # Check for name mentions and store profile automatically
        name_patterns = [
            r"my name is ([A-Za-z\s]+)",
            r"i am ([A-Za-z\s]+)",
            r"call me ([A-Za-z\s]+)",
            r"remember my name as ([A-Za-z\s]+)",
            r"i want you to remember my name as ([A-Za-z\s]+)"
        ]
        
        detected_name = None
        question_lower = request.question.lower()
        
        for pattern in name_patterns:
            match = re.search(pattern, question_lower)
            if match:
                detected_name = match.group(1).strip().title()
                break
        
        # If name detected, store it automatically
        if detected_name:
            try:
                # Add timestamp to username for uniqueness
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                unique_name = f"{detected_name}_{timestamp}"
                
                profile_text = f"User name: {unique_name}"
                profile_embedding = embeddings.embed_query(profile_text)
                profile_id = f"profile_{request.guest_id}"
                
                index.upsert(vectors=[{
                    "id": profile_id,
                    "values": profile_embedding,
                    "metadata": {
                        "type": "user_profile",
                        "user_id": request.guest_id,
                        "name": unique_name,
                        "display_name": detected_name,
                        "email": "",
                        "profile_text": profile_text
                    }
                }])
                
                logger.info(f"Auto-stored profile for user {request.guest_id}: {detected_name}")
                
                # Store this interaction in conversation memory
                memory = get_user_memory(request.guest_id)
                response_text = f"Nice to meet you, {detected_name}! I'll remember your name. How can I help you with your documents?"
                
                memory.save_context(
                    {"input": request.question},
                    {"output": response_text}
                )
                
                # Store in Pinecone for persistent memory
                store_conversation_in_pinecone(request.guest_id, request.question, response_text)
                
                return AnswerResponse(
                    answer=response_text,
                    sources=[]
                )
            except Exception as e:
                logger.error(f"Error auto-storing profile: {e}")
        
        # Get user profile for personalized responses
        user_profile = None
        try:
            profile_response = index.query(
                vector=[0] * 512,
                top_k=1,
                include_metadata=True,
                filter={"user_id": request.guest_id, "type": "user_profile"}
            )
            
            if profile_response.matches:
                user_profile = profile_response.matches[0].metadata
        except Exception as e:
            logger.error(f"Error retrieving user profile: {e}")
        
        # Get conversation memory for this user
        memory = get_user_memory(request.guest_id)
        
        # Get relevant conversation history
        relevant_history = get_relevant_conversation_history(request.guest_id, request.question)
        
        # Generate embedding for the question
        question_embedding = embeddings.embed_query(request.question)
        
        # Query Pinecone for relevant document chunks (user isolation)
        query_response = index.query(
            vector=question_embedding,
            top_k=5,  # Get top 5 most relevant chunks
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": {"$nin": ["user_profile", "conversation"]}}  # Exclude profiles and conversations
        )
        
        # Check if we have any document chunks
        document_chunks = [match for match in query_response.matches 
                          if match.metadata.get('type') not in ['user_profile', 'conversation']]
        
        if not document_chunks:
            user_name = user_profile.get("name", "") if user_profile else ""
            greeting = f"Hi {user_name}! " if user_name else "Hi! "
            
            # Check if this is a general greeting or conversation
            greeting_words = ['hi', 'hello', 'hey', 'good morning', 'good afternoon', 'good evening']
            if any(word in request.question.lower() for word in greeting_words):
                response_text = f"{greeting}How can I help you today? I can answer questions about your uploaded documents."
            else:
                response_text = f"{greeting}I couldn't find any relevant information in your uploaded documents to answer this question. Please make sure you have uploaded documents first, or ask me something else I can help with!"
            
            # Store conversation in memory
            memory.save_context(
                {"input": request.question},
                {"output": response_text}
            )
            
            # Store in Pinecone
            store_conversation_in_pinecone(request.guest_id, request.question, response_text)
            
            return AnswerResponse(
                answer=response_text,
                sources=[]
            )
        
        # Extract relevant chunks and prepare context
        relevant_chunks = []
        sources = []
        
        for match in document_chunks:
            metadata = match.metadata
            chunk_text = metadata.get('full_text', '')
            
            if chunk_text:
                relevant_chunks.append(chunk_text)
                
                # Prepare source information
                source_info = {
                    "document_name": metadata.get('document_name', 'Unknown'),
                    "chunk_index": metadata.get('chunk_index', 0),
                    "snippet": metadata.get('text', chunk_text[:200]),
                    "relevance_score": round(match.score, 3)
                }
                sources.append(source_info)
        
        # Create context from relevant chunks
        context = "\n\n".join(relevant_chunks)
        
        # Get conversation history from memory
        chat_history = memory.load_memory_variables({}).get('chat_history', [])
        
        # Format chat history for the prompt
        formatted_history = ""
        if chat_history:
            history_messages = []
            for msg in chat_history[-6:]:  # Last 3 exchanges (6 messages)
                if hasattr(msg, 'content'):
                    role = "Human" if isinstance(msg, HumanMessage) else "AI"
                    history_messages.append(f"{role}: {msg.content}")
            formatted_history = "\n".join(history_messages)
        
        # Add relevant past conversations if available
        if relevant_history:
            formatted_history += "\n\nRelevant past conversations:\n" + relevant_history
        
        logger.info(f"Generating answer for guest {request.guest_id} with {len(relevant_chunks)} relevant chunks")
        
        # Get user name for personalization
        user_name = user_profile.get("display_name", user_profile.get("name", "")) if user_profile else ""
        
        # Create enhanced prompt with conversation context
        enhanced_prompt_template = f"""
You are an intelligent and helpful AI assistant with access to the user's uploaded documents. You have a friendly, professional personality and excellent memory of previous conversations.

User Information:
- Name: {user_name if user_name else 'Not provided'}

Key Instructions:
1. Always provide helpful, accurate, and contextual responses based on the available documents
2. Remember previous conversations and refer to them when relevant
3. If information is missing or unclear in the documents, explicitly state what is available vs. unavailable
4. Be conversational and personable while maintaining professionalism
5. When greeting users, use their name if you know it
6. Provide specific details from documents rather than generic responses
7. If asked about something not in the documents, clearly state that and offer to help with what is available
8. For contact queries, be specific about what contact information IS available vs. what is missing
9. Never just say "I don't know" when you can provide partial information
10. **REPORT GENERATION**: When asked to create a report about specific topics from documents, provide comprehensive analysis with sections like: Executive Summary, Key Findings, Detailed Analysis, Recommendations, and Conclusion. Format reports professionally with clear headings and bullet points.
11. **EMAIL REPORTS**: When users request reports to be emailed, inform them about the 'Generate Report' feature that can create detailed reports and send them via email automatically.
12. If asked about people or topics not in the documents, clarify what information is available and suggest using the generate report feature for comprehensive analysis.
13. Always double-check the user's question for the correct spelling and context before providing information.

Document Context:
{{context}}

Conversation History:
{formatted_history}

Current Question: {request.question}

Provide a helpful and contextual response:"""
        
        # Create the enhanced prompt
        enhanced_qa_prompt = PromptTemplate.from_template(enhanced_prompt_template)
        enhanced_qa_chain = create_stuff_documents_chain(llm, enhanced_qa_prompt)
        
        # Create LangChain documents for the chain
        docs = [LangChainDocument(page_content=context)]
        
        # Run the QA chain
        result = enhanced_qa_chain.invoke({"context": docs})
        
        # Store the conversation in memory
        memory.save_context(
            {"input": request.question},
            {"output": result}
        )
        
        # Store in Pinecone for persistent memory
        store_conversation_in_pinecone(request.guest_id, request.question, result)
        
        return AnswerResponse(
            answer=result.strip(),
            sources=sources
        )
    
    except Exception as e:
        logger.error(f"Error answering question: {e}")
        raise HTTPException(status_code=500, detail=f"Error processing question: {str(e)}")

@app.post("/profile", response_model=UserProfileResponse)
async def store_user_profile(profile: UserProfile):
    """Store user profile information with embeddings"""
    try:
        # Add timestamp to username for uniqueness
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_name = f"{profile.name}_{timestamp}"
        
        # Create profile text for embedding
        profile_text = f"User name: {unique_name}"
        if profile.email:
            profile_text += f", Email: {profile.email}"
        
        # Generate embedding for profile
        profile_embedding = embeddings.embed_query(profile_text)
        
        # Store profile in Pinecone with special metadata
        profile_id = f"profile_{profile.guest_id}"
        
        index.upsert(vectors=[{
            "id": profile_id,
            "values": profile_embedding,
            "metadata": {
                "type": "user_profile",
                "user_id": profile.guest_id,
                "name": unique_name,
                "display_name": profile.name,  # Store original name for display
                "email": profile.email or "",
                "profile_text": profile_text
            }
        }])
        
        logger.info(f"Stored profile for user {profile.guest_id}: {unique_name}")
        
        return UserProfileResponse(
            message="Profile stored successfully",
            profile={
                "name": unique_name,
                "display_name": profile.name,
                "email": profile.email,
                "guest_id": profile.guest_id
            }
        )
    
    except Exception as e:
        logger.error(f"Error storing user profile: {e}")
        raise HTTPException(status_code=500, detail=f"Error storing profile: {str(e)}")

@app.get("/profile/{guest_id}")
async def get_user_profile(guest_id: str):
    """Retrieve user profile information"""
    try:
        # Query for user profile
        query_response = index.query(
            vector=[0] * 512,  # Dummy vector
            top_k=1,
            include_metadata=True,
            filter={"user_id": guest_id, "type": "user_profile"}
        )
        
        if query_response.matches:
            profile_data = query_response.matches[0].metadata
            return {
                "name": profile_data.get("display_name", profile_data.get("name", "")),  # Use display_name if available
                "full_name": profile_data.get("name", ""),  # Keep full unique name for backend use
                "email": profile_data.get("email", ""),
                "guest_id": guest_id
            }
        else:
            return {"name": "", "full_name": "", "email": "", "guest_id": guest_id}
    
    except Exception as e:
        logger.error(f"Error retrieving user profile: {e}")
        raise HTTPException(status_code=500, detail=f"Error retrieving profile: {str(e)}")

@app.post("/generate-report")
async def generate_ai_report(request: GenerateReportRequest):
    """Generate AI-powered comprehensive report from document content"""
    try:
        # Get user profile
        profile_response = index.query(
            vector=[0] * 512,
            top_k=1,
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": "user_profile"}
        )
        
        user_profile = None
        if profile_response.matches:
            user_profile = profile_response.matches[0].metadata
        
        # Generate embedding for the report topic
        topic_embedding = embeddings.embed_query(request.report_topic)
        
        # Query Pinecone for ALL relevant document chunks
        query_response = index.query(
            vector=topic_embedding,
            top_k=20,  # Get more chunks for comprehensive report
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": {"$nin": ["user_profile", "conversation"]}}
        )
        
        if not query_response.matches:
            return {
                "success": False,
                "message": "No relevant documents found to generate a report.",
                "report": None
            }
        
        # Extract all relevant content
        all_content = []
        for match in query_response.matches:
            metadata = match.metadata
            chunk_text = metadata.get('full_text', '')
            if chunk_text:
                all_content.append(chunk_text)
        
        # Create comprehensive context
        comprehensive_context = "\n\n".join(all_content)
        
        # Get user name for personalization
        user_name = user_profile.get("display_name", user_profile.get("name", "")) if user_profile else ""
        
        # Create specialized report generation prompt
        report_prompt_template = f"""
You are an expert report writer and data analyst. Create a comprehensive, professional report based on the provided document content.

User Information:
- Name: {user_name if user_name else 'Not provided'}
- Report Topic: {request.report_topic}

Instructions:
1. Analyze ALL the provided document content thoroughly
2. Create a well-structured, professional report with the following sections:
   - **EXECUTIVE SUMMARY** (2-3 sentences overview)
   - **DETAILED FINDINGS** (organized by relevant topics/categories)
   - **KEY DATA & STATISTICS** (any numbers, dates, quantities found)
   - **CONTACT INFORMATION** (any emails, phones, addresses found)
   - **RECOMMENDATIONS/CONCLUSIONS** (based on the analysis)
3. Use professional formatting with clear headers and bullet points
4. Include specific details and quotes from the documents
5. If the topic doesn't match the document content, explain what IS available instead
6. Make the report comprehensive and actionable

Document Content:
{comprehensive_context}

Generate a comprehensive professional report:"""
        
        # Create the report generation prompt
        report_qa_prompt = PromptTemplate.from_template(report_prompt_template)
        report_qa_chain = create_stuff_documents_chain(llm, report_qa_prompt)
        
        # Create LangChain documents for the chain
        docs = [LangChainDocument(page_content=comprehensive_context)]
        
        # Generate the report
        report_result = report_qa_chain.invoke({"context": docs})
        
        # If email is requested, send it
        email_sent = False
        if request.include_email and request.email_address:
            if EMAIL_ADDRESS and EMAIL_PASSWORD:
                subject = f"AI Generated Report: {request.report_topic}"
                email_content = f"""Hello {user_name if user_name else 'User'},

Here is your requested AI-generated report:

{report_result}

---
Generated by AI Document Assistant
Report ID: {request.guest_id}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}

This is an automated report. Please do not reply to this email."""
                
                email_sent = send_email(request.email_address, subject, email_content)
        
        return {
            "success": True,
            "message": "Report generated successfully!" + (" Email sent!" if email_sent else ""),
            "report": report_result,
            "email_sent": email_sent,
            "chunks_analyzed": len(all_content)
        }
        
    except Exception as e:
        logger.error(f"Error generating AI report: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")

@app.post("/send-report")
async def send_email_report(request: EmailReportRequest):
    """Send email report to user"""
    try:
        # Get user profile
        profile_response = index.query(
            vector=[0] * 512,
            top_k=1,
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": "user_profile"}
        )
        
        user_name = "User"
        if profile_response.matches:
            metadata = profile_response.matches[0].metadata
            user_name = metadata.get("display_name", metadata.get("name", "User"))
        
        # Get user's document count and details
        doc_response = index.query(
            vector=[0] * 512,
            top_k=100,
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": {"$ne": "user_profile"}}
        )
        
        doc_count = len(doc_response.matches)
        
        # Get unique document names
        doc_names = set()
        for match in doc_response.matches:
            doc_name = match.metadata.get('document_name', 'Unknown')
            if doc_name != 'Unknown':
                doc_names.add(doc_name)
        
        # Get recent conversations for summary
        conversation_response = index.query(
            vector=[0] * 512,
            top_k=10,
            include_metadata=True,
            filter={"user_id": request.guest_id, "type": "conversation"}
        )
        
        # Extract conversation insights
        recent_questions = []
        for match in conversation_response.matches:
            question = match.metadata.get('question', '')
            if question and len(recent_questions) < 5:
                recent_questions.append(question)
        
        # Get document insights
        doc_insights = []
        if doc_names:
            for doc_name in sorted(doc_names):
                doc_chunks = [m for m in doc_response.matches if m.metadata.get('document_name') == doc_name]
                doc_insights.append(f"• {doc_name}: {len(doc_chunks)} chunks processed")
        
        # Create enhanced detailed report
        report_content = f"""AI Document Assistant - Enhanced Activity Report

Hello {user_name},

Here's your comprehensive activity summary:

📊 SUMMARY
• Guest ID: {request.guest_id}
• Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
• Total Document Chunks: {doc_count}
• Unique Documents: {len(doc_names)}
• Recent Conversations: {len(recent_questions)}

📄 UPLOADED DOCUMENTS
{chr(10).join(doc_insights) if doc_insights else '• No documents uploaded yet'}

💬 RECENT ACTIVITY
{chr(10).join([f'• "{q[:80]}..."' if len(q) > 80 else f'• "{q}"' for q in recent_questions]) if recent_questions else '• No recent questions'}

🤖 SYSTEM STATUS
• Profile Status: {'✅ Active' if profile_response.matches else '❌ Not Set'}
• Data Isolation: ✅ Secure (Guest ID based)
• Vector Storage: ✅ Active in Pinecone
• AI Memory: ✅ Conversation history maintained

📈 INSIGHTS
• Most active document: {max(doc_names, key=lambda x: len([m for m in doc_response.matches if m.metadata.get('document_name') == x])) if doc_names else 'None'}
• Total interactions: {len(recent_questions) + doc_count}
• Account activity: {'High' if len(recent_questions) > 3 else 'Moderate' if len(recent_questions) > 0 else 'Low'}

Thank you for using AI Document Assistant!

Best regards,
AI Document Assistant Team

---
This is an automated report. Please do not reply to this email."""
        
        # Try to send email if configured
        email_sent = False
        if EMAIL_ADDRESS and EMAIL_PASSWORD:
            subject = f"AI Document Assistant Report - {user_name}"
            email_sent = send_email(request.email, subject, report_content)
        
        # Log the report regardless
        logger.info(f"Email report generated for {request.guest_id} ({user_name})")
        
        if email_sent:
            return {
                "message": f"Report sent successfully to {request.email}!",
                "email_sent": True,
                "report_preview": report_content
            }
        else:
            return {
                "message": f"Report generated for {user_name}. Email configuration needed to send emails.",
                "email_sent": False,
                "report_preview": report_content,
                "setup_instructions": "To enable email sending, add EMAIL_ADDRESS and EMAIL_PASSWORD to your .env file"
            }
    
    except Exception as e:
        logger.error(f"Error generating email report: {e}")
        raise HTTPException(status_code=500, detail=f"Error generating report: {str(e)}")

@app.delete("/documents/{guest_id}/{document_name}")
async def delete_document_by_path(guest_id: str, document_name: str):
    """Delete a specific document for a guest user via URL path"""
    try:
        # Query vectors for this specific document
        query_response = index.query(
            vector=[0] * 512,
            top_k=1000,
            include_metadata=True,
            filter={
                "user_id": guest_id,
                "document_name": document_name
            }
        )
        
        # Also query for related conversation chunks that might reference this document
        conversation_response = index.query(
            vector=[0] * 512,
            top_k=1000,
            include_metadata=True,
            filter={
                "user_id": guest_id,
                "type": "conversation"
            }
        )
        
        document_chunks_deleted = 0
        conversation_chunks_deleted = 0
        
        if query_response.matches:
            # Extract IDs to delete for document chunks
            document_ids_to_delete = [match.id for match in query_response.matches]
            document_chunks_deleted = len(document_ids_to_delete)
            
            # Delete document vectors from Pinecone
            index.delete(ids=document_ids_to_delete)
            
        # Delete related conversation chunks that mention this document
        if conversation_response.matches:
            conversation_ids_to_delete = []
            for match in conversation_response.matches:
                # Check if conversation mentions the document name
                question = match.metadata.get('question', '').lower()
                answer = match.metadata.get('answer', '').lower()
                if document_name.lower() in question or document_name.lower() in answer:
                    conversation_ids_to_delete.append(match.id)
            
            if conversation_ids_to_delete:
                conversation_chunks_deleted = len(conversation_ids_to_delete)
                index.delete(ids=conversation_ids_to_delete)
        
        if document_chunks_deleted > 0 or conversation_chunks_deleted > 0:
            logger.info(f"Deleted document '{document_name}' ({document_chunks_deleted} document chunks, {conversation_chunks_deleted} conversation chunks) for guest {guest_id}")
            return {
                "message": f"Successfully deleted document '{document_name}' with {document_chunks_deleted} document chunks and {conversation_chunks_deleted} related conversation chunks",
                "document_chunks_deleted": document_chunks_deleted,
                "conversation_chunks_deleted": conversation_chunks_deleted
            }
        else:
            raise HTTPException(
                status_code=404,
                detail=f"Document '{document_name}' not found for this user"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.delete("/document")
async def delete_document(request: DeleteDocumentRequest):
    """Delete a specific document for a guest user via JSON body"""
    try:
        # Query vectors for this specific document
        query_response = index.query(
            vector=[0] * 512,
            top_k=1000,
            include_metadata=True,
            filter={
                "user_id": request.guest_id,
                "document_name": request.document_name
            }
        )
        
        if query_response.matches:
            # Extract IDs to delete
            ids_to_delete = [match.id for match in query_response.matches]
            
            # Delete vectors from Pinecone
            index.delete(ids=ids_to_delete)
            
            logger.info(f"Deleted document '{request.document_name}' ({len(ids_to_delete)} chunks) for guest {request.guest_id}")
            return {
                "message": f"Successfully deleted document '{request.document_name}' with {len(ids_to_delete)} chunks",
                "chunks_deleted": len(ids_to_delete)
            }
        else:
            raise HTTPException(
                status_code=404,
                detail=f"Document '{request.document_name}' not found for this user"
            )
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error deleting document: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.delete("/conversations/{guest_id}")
async def delete_conversations(guest_id: str):
    """Delete all conversation chunks for a specific guest user"""
    try:
        # Query conversation vectors for this user
        query_response = index.query(
            vector=[0] * 512,
            top_k=10000,
            include_metadata=True,
            filter={
                "user_id": guest_id,
                "type": "conversation"
            }
        )
        
        if query_response.matches:
            # Extract IDs to delete
            ids_to_delete = [match.id for match in query_response.matches]
            
            # Delete conversation vectors from Pinecone
            index.delete(ids=ids_to_delete)
            
            logger.info(f"Deleted {len(ids_to_delete)} conversation chunks for guest {guest_id}")
            return {
                "message": f"Successfully deleted all conversations",
                "conversation_chunks_deleted": len(ids_to_delete)
            }
        else:
            return {
                "message": "No conversations found for this user",
                "conversation_chunks_deleted": 0
            }
            
    except Exception as e:
        logger.error(f"Error deleting conversations: {e}")
        raise HTTPException(status_code=500, detail=f"Error deleting conversations: {str(e)}")

@app.delete("/clear/{guest_id}")
async def clear_user_data(guest_id: str):
    """Clear all data for a specific guest user including documents, conversations, and other data types"""
    try:
        # Query all vectors for this user
        query_response = index.query(
            vector=[0] * 512,
            top_k=10000,  # Large number to get all vectors
            include_metadata=True,
            filter={"user_id": guest_id}
        )
        
        if query_response.matches:
            # Categorize chunks by type
            document_chunks = []
            conversation_chunks = []
            other_chunks = []
            
            for match in query_response.matches:
                chunk_type = match.metadata.get('type', 'document')
                if chunk_type == 'conversation':
                    conversation_chunks.append(match.id)
                elif chunk_type == 'document' or 'document_name' in match.metadata:
                    document_chunks.append(match.id)
                else:
                    other_chunks.append(match.id)
            
            # Extract all IDs to delete
            ids_to_delete = [match.id for match in query_response.matches]
            
            # Delete all vectors from Pinecone
            index.delete(ids=ids_to_delete)
            
            logger.info(f"Cleared all data for guest {guest_id}: {len(document_chunks)} document chunks, {len(conversation_chunks)} conversation chunks, {len(other_chunks)} other chunks")
            return {
                "message": f"Successfully cleared all data for user",
                "total_chunks_deleted": len(ids_to_delete),
                "document_chunks_deleted": len(document_chunks),
                "conversation_chunks_deleted": len(conversation_chunks),
                "other_chunks_deleted": len(other_chunks)
            }
        else:
            return {
                "message": "No data found for this user",
                "total_chunks_deleted": 0,
                "document_chunks_deleted": 0,
                "conversation_chunks_deleted": 0,
                "other_chunks_deleted": 0
            }
            
    except Exception as e:
        logger.error(f"Error clearing user data: {e}")
        raise HTTPException(status_code=500, detail=f"Error clearing user data: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)